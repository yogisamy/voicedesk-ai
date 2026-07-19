"""
Generates the Assignment II submission Word document for VoiceDesk AI.
Run: python3 generate_report.py
Output: groupid.docx (rename to your actual group ID before submitting)
"""

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SCREENSHOTS_DIR = "."
SCREENSHOTS = {
    "support_pipeline": "Screenshot 2026-07-19 at 9.27.03 AM.png",
    "dashboard": "Screenshot 2026-07-19 at 9.27.21 AM.png",
    "mlflow_runs": "Screenshot 2026-07-19 at 9.28.01 AM.png",
    "mlflow_metrics": "Screenshot 2026-07-19 at 9.28.16 AM.png",
    "mlflow_charts": "Screenshot 2026-07-19 at 9.28.33 AM.png",
}


# ── helpers ──────────────────────────────────────────────────────────────────

def set_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return h


def add_para(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    return p


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light List Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # header row
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
    # data rows
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
    if col_widths:
        for row in table.rows:
            for ci, w in enumerate(col_widths):
                row.cells[ci].width = Inches(w)
    doc.add_paragraph()
    return table


def add_screenshot(doc, key, caption, width=Inches(5.5)):
    path = os.path.join(SCREENSHOTS_DIR, SCREENSHOTS.get(key, ""))
    if os.path.exists(path):
        doc.add_picture(path, width=width)
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
        c = doc.add_paragraph(caption)
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in c.runs:
            run.italic = True
            run.font.size = Pt(9)
    else:
        doc.add_paragraph(f"[Screenshot: {caption}]")
    doc.add_paragraph()


# ── main document ─────────────────────────────────────────────────────────────

def build():
    doc = Document()

    # ── Title page ────────────────────────────────────────────────────────────
    title = doc.add_heading("VoiceDesk AI", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph("API-driven AI Customer Support Assistant")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in sub.runs:
        r.font.size = Pt(14)

    course = doc.add_paragraph("CCZG506 — API-driven Cloud Native Solutions | Assignment II")
    course.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # ── Group Details ─────────────────────────────────────────────────────────
    set_heading(doc, "Group Details", 1)

    add_table(doc,
        headers=["Sl. No", "BITS ID", "Name", "Contribution (Qualitative)", "% Contribution"],
        rows=[
            ["1", "<BITS ID>", "<Name>", "<Describe contribution>", "<e.g. 20%>"],
            ["2", "<BITS ID>", "<Name>", "<Describe contribution>", "<e.g. 20%>"],
            ["3", "<BITS ID>", "<Name>", "<Describe contribution>", "<e.g. 20%>"],
            ["4", "<BITS ID>", "<Name>", "<Describe contribution>", "<e.g. 20%>"],
            ["5", "<BITS ID>", "<Name>", "<Describe contribution>", "<e.g. 20%>"],
        ],
        col_widths=[0.4, 1.1, 1.4, 2.8, 1.0],
    )
    add_para(doc, "Note: Fill in the group details above before submission.", italic=True)

    doc.add_page_break()

    # ── 1. Project Overview ───────────────────────────────────────────────────
    set_heading(doc, "1. Project Overview", 1)
    add_para(doc,
        "VoiceDesk AI is an end-to-end, API-driven customer support assistant built "
        "as a single cohesive Gradio web application. A customer records or types a "
        "complaint; the system transcribes it, detects sentiment, classifies intent "
        "(using both a zero-shot model and our own fine-tuned model), extracts named "
        "entities, summarises the ticket for a human agent, drafts an empathetic "
        "AI-generated reply, and speaks it back — all while logging LLMOps metrics "
        "for every API call."
    )

    add_table(doc,
        headers=["Assignment Requirement", "How Satisfied"],
        rows=[
            ["Domain", "Customer Service"],
            ["Two categories", "NLP + Speech Recognition"],
            ["≥5 sub-tasks", "7 sub-tasks implemented (see Section 3)"],
            ["LLM/SLM via APIs", "Hugging Face Inference API (Whisper, DistilBERT, BART, Qwen2.5) + gTTS"],
            ["Cohesive unified objective", "One end-to-end support-ticket pipeline"],
            ["Interactive & demonstrable", "Gradio web application"],
            ["LLMOps, ≥5 metrics", "6 metrics + in-app dashboard + MLflow experiment tracking"],
            ["Fine-tune on same-domain dataset", "DistilBERT fine-tuned on Bitext Customer Support dataset"],
        ],
        col_widths=[2.5, 4.2],
    )

    # ── 2. Architecture ───────────────────────────────────────────────────────
    set_heading(doc, "2. Architecture", 1)

    set_heading(doc, "2.1 High-Level Design (HLD)", 2)
    add_para(doc,
        "The application has three main layers:"
    )
    add_bullet(doc, "User (Browser) — inputs a voice recording or typed complaint; receives results and spoken reply.")
    add_bullet(doc, "Gradio Web App (app.py) — orchestrates the full pipeline, hosts the UI, logs metrics, and serves the LLMOps Dashboard.")
    add_bullet(doc, "Cloud APIs — Hugging Face Inference API (ASR, NLP models) and gTTS for speech synthesis.")
    add_para(doc,
        "Metrics are written to metrics_log.csv (CSV) and simultaneously to an MLflow "
        "tracking store (SQLite) so both a built-in dashboard and the MLflow UI can "
        "visualise them."
    )

    set_heading(doc, "2.2 Low-Level Design (LLD) — Pipeline Sequence", 2)
    add_para(doc,
        "When the user clicks ▶ Run Full AI Pipeline, the run_pipeline() generator "
        "executes the following steps in order, yielding live results to the UI after each:"
    )
    steps = [
        ("Step 1", "Speech-to-Text", "Audio → Whisper-large-v3 (HF) → transcript"),
        ("Step 2", "Sentiment Analysis", "Transcript → DistilBERT-SST-2 (HF) → POSITIVE/NEGATIVE + confidence"),
        ("Step 3", "Intent Classification", "Transcript → BART-large-MNLI (zero-shot) AND → fine-tuned DistilBERT (local)"),
        ("Step 4", "Named Entity Recognition", "Transcript → BERT-base-NER (HF) → PER/ORG/LOC entities"),
        ("Step 5", "Summarization", "Transcript → BART-large-CNN (HF) → ticket summary + ROUGE-1 F1"),
        ("Step 6", "Response Generation", "Transcript + sentiment + intent → Qwen2.5-7B (HF) → empathetic reply"),
        ("Step 7", "Text-to-Speech", "Reply → gTTS API → MP3 audio played in browser"),
    ]
    add_table(doc,
        headers=["Step", "Sub-task", "Data Flow"],
        rows=steps,
        col_widths=[0.6, 1.8, 4.3],
    )
    add_para(doc,
        "Every step calls log_metric() which appends a row to metrics_log.csv and "
        "(if MLflow is installed) mirrors the metric to the active MLflow run."
    )

    # ── 3. Sub-tasks ─────────────────────────────────────────────────────────
    set_heading(doc, "3. Sub-tasks Implemented", 1)

    subtasks = [
        ("1", "Speech-to-Text", "Speech Recognition",
         "openai/whisper-large-v3 (HF Inference API)",
         "Converts the customer's recorded/uploaded voice complaint into text. "
         "The transcript feeds every downstream NLP sub-task. "
         "Metrics: latency, word count, character count."),
        ("2", "Text-to-Speech", "Speech Recognition",
         "gTTS (Google Text-to-Speech API)",
         "Speaks the AI-generated reply back to the customer (first 500 chars), "
         "saving an MP3 that plays in the browser. "
         "Metrics: latency, word count."),
        ("3", "Sentiment Analysis", "NLP",
         "distilbert-base-uncased-finetuned-sst-2-english (HF)",
         "Classifies the complaint as POSITIVE/NEGATIVE. The detected sentiment "
         "is injected into the reply-generation prompt so the LLM matches its "
         "tone to the customer's mood. "
         "Metrics: latency, confidence."),
        ("4a", "Intent Classification — Zero-shot", "NLP",
         "facebook/bart-large-mnli (HF)",
         "Scores the complaint against 7 candidate intents (refund request, "
         "cancel order, delivery problem, payment issue, product complaint, "
         "account help, general enquiry) with no training needed. "
         "Metrics: latency, confidence."),
        ("4b", "Intent Classification — Fine-tuned", "NLP",
         "DistilBERT fine-tuned on Bitext Customer Support dataset (local ./intent-model)",
         "Our trained model (27 real-world support intents) runs in parallel with "
         "the zero-shot baseline. Both predictions are shown side-by-side so the "
         "quality difference is directly visible. "
         "Metrics: latency, confidence, accuracy, macro-F1."),
        ("5", "Named Entity Recognition", "NLP",
         "dslim/bert-base-NER (HF)",
         "Extracts people (PER), organisations (ORG), and locations (LOC) from the "
         "complaint — giving the human agent structured ticket fields. "
         "Metrics: latency, average entity confidence, entity count."),
        ("6", "Summarization", "NLP",
         "facebook/bart-large-cnn (HF)",
         "Condenses long complaints into a short ticket summary for the human "
         "agent. Summary quality is scored automatically with ROUGE-1 F1. "
         "Metrics: latency, ROUGE-1 F1, token count."),
        ("7", "Response Generation", "NLP",
         "Qwen/Qwen2.5-7B-Instruct (HF router, OpenAI-compatible chat API)",
         "Drafts a short (≤120 words), empathetic, professional reply with a "
         "concrete next step. The prompt is grounded with sentiment + detected intent. "
         "Metrics: latency, total tokens, estimated cost (USD)."),
    ]

    for num, name, cat, model, desc in subtasks:
        set_heading(doc, f"3.{num}  {name}  [{cat}]", 2)
        add_para(doc, f"Model: {model}", bold=True)
        add_para(doc, desc)

    # ── 4. LLMOps Metrics ────────────────────────────────────────────────────
    set_heading(doc, "4. LLMOps — Metrics Measured", 1)

    add_para(doc,
        "All metrics are logged in real-time to metrics_log.csv and mirrored to an "
        "MLflow experiment. Six distinct metrics are tracked:"
    )

    add_table(doc,
        headers=["#", "Metric", "Captured Where", "Purpose"],
        rows=[
            ["1", "Latency per sub-task (seconds)", "Every API call", "Identify slowest pipeline steps"],
            ["2", "Model confidence (0–1)", "Sentiment, Intent, NER", "Monitor prediction quality over time"],
            ["3", "Intent model quality: Accuracy & Macro-F1", "Fine-tuning script + live side-by-side comparison", "Quantify improvement of fine-tuning vs zero-shot"],
            ["4", "Summary quality (ROUGE-1 F1)", "Summarization sub-task", "Automated quality signal for summaries"],
            ["5", "Token usage & estimated cost (USD)", "LLM response generation", "Cost visibility for production LLMOps"],
            ["6", "User satisfaction (👍 / 👎 feedback)", "In-app feedback widget", "Human evaluation of reply quality"],
        ],
        col_widths=[0.3, 1.7, 1.9, 2.8],
    )

    set_heading(doc, "4.1 MLflow Experiment Tracking", 2)
    add_para(doc,
        "Every metric is also mirrored to MLflow (industry-standard LLMOps tooling):"
    )
    add_bullet(doc,
        "voicedesk-ai experiment — one MLflow run per pipeline execution, with "
        "per-sub-task latency/confidence/token metrics, input-mode and "
        "transcript-length parameters, and one run per thumbs-up/thumbs-down feedback event.")
    add_bullet(doc,
        "voicedesk-intent-finetune experiment — one run per fine-tuning job, "
        "logging hyperparameters (base model, dataset, epochs, learning rate, "
        "batch size) and final test accuracy / macro-F1 / loss.")
    add_para(doc, "Browse both at http://localhost:5000 by running (from the project folder):", italic=True)
    doc.add_paragraph("    mlflow ui --backend-store-uri sqlite:///mlflow.db").style = "No Spacing"
    doc.add_paragraph()

    # ── 5. Fine-tuning ────────────────────────────────────────────────────────
    set_heading(doc, "5. Fine-tuning (Requirement #8)", 1)

    add_table(doc,
        headers=["Parameter", "Value"],
        rows=[
            ["Base model", "distilbert-base-uncased (66M parameters — SLM)"],
            ["Dataset", "Bitext Customer Support (27k utterances, 27 real intents)"],
            ["Training samples", "8,000 (subsampled for speed; full dataset can be used)"],
            ["Epochs", "2"],
            ["Learning rate", "2e-5"],
            ["Train batch size", "32"],
            ["Eval batch size", "64"],
            ["Hardware", "Apple M-series (MPS) — ~100s locally"],
            ["Test Accuracy", "97.07%"],
            ["Test Macro-F1", "97.07%"],
        ],
        col_widths=[2.2, 4.5],
    )

    add_para(doc,
        "The fine-tuned model is saved to ./intent-model and loaded by app.py at "
        "startup. The UI shows the fine-tuned prediction alongside the zero-shot "
        "BART prediction so the quality improvement is directly visible. "
        "The training run's hyperparameters and final metrics are logged to the "
        "voicedesk-intent-finetune MLflow experiment."
    )

    add_para(doc,
        "Why DistilBERT (SLM)? It is 40% smaller and 60% faster than BERT while "
        "retaining ~97% of BERT's performance — ideal for low-latency "
        "intent classification in a real-time support pipeline.",
        italic=True
    )

    # ── 6. Application Screenshots ────────────────────────────────────────────
    set_heading(doc, "6. Application Screenshots", 1)

    set_heading(doc, "6.1 Support Pipeline — Full Pipeline Output", 2)
    add_para(doc,
        "The screenshot below shows a complete pipeline run for the sample complaint: "
        "\"Hi, this is Priya. My order 4521 from Chennai arrived damaged and I want a refund.\""
    )
    add_para(doc, "Visible outputs:")
    add_bullet(doc, "Transcript (Whisper ASR)")
    add_bullet(doc, "Sentiment: NEGATIVE (99.87% confidence)")
    add_bullet(doc, "Zero-shot intent: refund request (BART-MNLI)")
    add_bullet(doc, "Fine-tuned intent: track_refund / get_refund (DistilBERT)")
    add_bullet(doc, "Named entities: Priya=PER, Chennai=LOC (highlighted)")
    add_bullet(doc, "Ticket summary")
    add_bullet(doc, "AI-generated empathetic reply (Qwen2.5-7B)")
    add_bullet(doc, "Spoken reply audio (gTTS)")
    add_screenshot(doc, "support_pipeline",
                   "Figure 1: VoiceDesk AI — Support Pipeline tab with full pipeline output")

    set_heading(doc, "6.2 LLMOps Dashboard", 2)
    add_para(doc,
        "The Dashboard tab aggregates all data from metrics_log.csv in real-time. "
        "It shows:"
    )
    add_bullet(doc, "KPI bar: total API calls, average latency, total tokens, estimated cost, user satisfaction")
    add_bullet(doc, "Bar chart: average latency per sub-task (summarization ~7.8s; fine-tuned intent ~0.25s)")
    add_bullet(doc, "Bar chart: average model confidence per sub-task (sentiment 0.99, NER 0.84)")
    add_bullet(doc, "Per-sub-task metrics table from metrics_log.csv")
    add_screenshot(doc, "dashboard",
                   "Figure 2: LLMOps Dashboard — aggregated metrics from all pipeline runs")

    set_heading(doc, "6.3 MLflow Experiment — Run List", 2)
    add_para(doc,
        "MLflow captures every pipeline execution as a distinct run. The screenshot "
        "shows three completed runs in the voicedesk-ai experiment, each with "
        "duration and source file recorded."
    )
    add_screenshot(doc, "mlflow_runs",
                   "Figure 3: MLflow — voicedesk-ai experiment run list")

    set_heading(doc, "6.4 MLflow — Per-Run Metrics Detail", 2)
    add_para(doc,
        "Drilling into a single run shows all 18 metrics logged for that pipeline "
        "execution: per-sub-task latency, confidence, token counts, and estimated costs."
    )
    add_screenshot(doc, "mlflow_metrics",
                   "Figure 4: MLflow — per-run metrics overview (pipeline-20260719-092612)")

    set_heading(doc, "6.5 MLflow — Metric Charts", 2)
    add_para(doc,
        "MLflow's model metrics tab plots each metric as a chart across runs, "
        "enabling trend analysis over time."
    )
    add_screenshot(doc, "mlflow_charts",
                   "Figure 5: MLflow — metric charts for intent confidence and latency")

    # ── 7. Setup & Run Instructions ───────────────────────────────────────────
    set_heading(doc, "7. Setup & Run Instructions", 1)

    set_heading(doc, "7.1 Running the Application", 2)
    steps_run = [
        ("1", "Get a free Hugging Face token", "https://huggingface.co/settings/tokens (Read access)"),
        ("2", "Install dependencies", "pip3 install -r requirements.txt"),
        ("3", "Set the token", "export HF_TOKEN=hf_xxxxxxxxxxxx"),
        ("4", "Launch the app", "python3 app.py"),
        ("5", "Open the URL", "Gradio prints a local URL (e.g. http://127.0.0.1:7860)"),
        ("6", "(Optional) View MLflow", "mlflow ui --backend-store-uri sqlite:///mlflow.db  →  http://127.0.0.1:5000"),
    ]
    add_table(doc,
        headers=["Step", "Action", "Command / URL"],
        rows=steps_run,
        col_widths=[0.4, 1.8, 4.5],
    )

    set_heading(doc, "7.2 Regenerating the Fine-tuned Model", 2)
    add_para(doc,
        "The fine-tuned model is stored in ./intent-model next to app.py. "
        "To retrain it (e.g. after a clean clone):"
    )
    add_bullet(doc, "pip3 install transformers datasets evaluate accelerate scikit-learn")
    add_bullet(doc, "python3 finetune_intent.py   (runs in ~100s on Apple Silicon MPS or ~10 min on a free Colab T4 GPU)")
    add_bullet(doc, "The script prints final Accuracy and Macro-F1 and saves the model to ./intent-model")

    set_heading(doc, "7.3 Demo Script (Viva)", 2)
    add_bullet(doc, "Record: \"Hi, this is Priya. My order 4521 from Chennai arrived damaged and I want a refund.\"")
    add_bullet(doc, "Show: transcript → sentiment (NEGATIVE) → both intent models → NER highlights (Priya, Chennai) → summary → reply → TTS audio")
    add_bullet(doc, "Rate the reply 👍 → open Dashboard tab → Refresh → show live metrics")
    add_bullet(doc, "Explain fine-tuning: base model choice (DistilBERT SLM), dataset (Bitext), accuracy/F1, and why SLM was chosen over a larger model")

    # ── 8. File Structure ─────────────────────────────────────────────────────
    set_heading(doc, "8. Project File Structure", 1)

    add_table(doc,
        headers=["File / Folder", "Purpose"],
        rows=[
            ["app.py", "Main Gradio web application — 7 sub-tasks, LLMOps logging, Dashboard"],
            ["finetune_intent.py", "DistilBERT fine-tuning script on Bitext Customer Support dataset"],
            ["requirements.txt", "Python dependencies"],
            ["intent-model/", "Saved fine-tuned DistilBERT (produced by finetune_intent.py)"],
            ["metrics_log.csv", "CSV log of every API call (latency, confidence, tokens, cost)"],
            ["mlflow.db", "MLflow SQLite store — browse with `mlflow ui --backend-store-uri sqlite:///mlflow.db`"],
            ["finetune_run.log", "Console output from the fine-tuning run (accuracy/F1 evidence)"],
            ["README.md", "Full project documentation with HLD/LLD Mermaid diagrams"],
        ],
        col_widths=[1.8, 4.9],
    )

    doc.save("groupid.docx")
    print("groupid.docx generated successfully.")
    print("Rename it to your actual group ID before uploading to the portal.")


if __name__ == "__main__":
    build()
