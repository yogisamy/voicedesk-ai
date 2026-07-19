"""
Generates a detailed code-explanation Word document for the VoiceDesk AI project.
Run: python3 generate_code_explainer.py
Output: VoiceDesk_AI_Code_Explainer.docx
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── helpers ──────────────────────────────────────────────────────────────────

def heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return h


def para(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p


def bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    return p


def code_block(doc, text):
    """Monospace grey paragraph for code snippets."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    # light grey background via XML shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F2F2F2")
    pPr.append(shd)
    return p


def table(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light List Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
    for ri, row in enumerate(rows):
        cells = t.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
    if col_widths:
        for row in t.rows:
            for ci, w in enumerate(col_widths):
                row.cells[ci].width = Inches(w)
    doc.add_paragraph()
    return t


# ── document ─────────────────────────────────────────────────────────────────

def build():
    doc = Document()

    # Title
    t = doc.add_heading("VoiceDesk AI — Complete Code Explainer", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("For team members | CCZG506 Assignment II")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # ── Overview ──────────────────────────────────────────────────────────────
    heading(doc, "1. Project Overview", 1)
    para(doc,
        "VoiceDesk AI is a two-file Python project:"
    )
    bullet(doc, "app.py — the Gradio web application (main deliverable, ~430 lines)")
    bullet(doc, "finetune_intent.py — standalone fine-tuning script (~150 lines)")
    para(doc,
        "Together they implement an end-to-end AI customer-support pipeline covering "
        "7 sub-tasks across NLP and Speech Recognition, driven entirely by cloud APIs "
        "(Hugging Face Inference API + gTTS). LLMOps metrics are logged to a CSV "
        "file and to MLflow on every run."
    )

    # ── How to Run ────────────────────────────────────────────────────────────
    heading(doc, "2. How to Run (Quick Reference)", 1)

    heading(doc, "2.1  Getting a Hugging Face Token", 2)
    para(doc,
        "All cloud models (Whisper, DistilBERT, BART, Qwen2.5) are served through "
        "the Hugging Face Inference API. A FREE account token is required:"
    )
    bullet(doc, "Go to https://huggingface.co/settings/tokens")
    bullet(doc, "Click 'New token' → choose 'Read' access → copy the hf_xxx... string")
    bullet(doc, "Each team member creates their own token (it's personal, like a password)")
    para(doc, "Why is a token needed? HF rate-limits unauthenticated requests to ~10/hour. "
         "A free token raises this to a comfortable level for demos.", italic=True)

    heading(doc, "2.2  First-Time Setup", 2)
    code_block(doc, "pip3 install -r requirements.txt")
    code_block(doc, "export HF_TOKEN=hf_your_token_here   # Mac/Linux")
    code_block(doc, "set HF_TOKEN=hf_your_token_here      # Windows")
    code_block(doc, "python3 app.py")
    para(doc, "Gradio prints a local URL (e.g. http://127.0.0.1:7860). Open it in a browser.")

    heading(doc, "2.3  Fine-tuning (One-Time, Already Done)", 2)
    para(doc,
        "The fine-tuned model lives in ./intent-model next to app.py. "
        "If that folder is missing (e.g. after a fresh clone), regenerate it with:"
    )
    code_block(doc, "python3 finetune_intent.py    # ~100 seconds on Apple Silicon")
    para(doc, "The script downloads the Bitext dataset from HuggingFace Hub automatically — "
         "no JSONL file or manual data download is needed.", italic=True)

    heading(doc, "2.4  (Optional) MLflow Dashboard", 2)
    code_block(doc, "mlflow ui --backend-store-uri sqlite:///mlflow.db")
    para(doc, "Then open http://127.0.0.1:5000 in a browser to browse experiment runs.")

    doc.add_page_break()

    # ── app.py deep dive ──────────────────────────────────────────────────────
    heading(doc, "3. app.py — Deep Dive", 1)
    para(doc,
        "app.py is organised into clearly labelled sections. We'll go through each "
        "one in the order they appear in the file."
    )

    # 3.1 Imports
    heading(doc, "3.1  Imports (Lines 1–35)", 2)
    table(doc,
        headers=["Import", "What it does"],
        rows=[
            ["gradio", "Builds the interactive web UI with zero HTML/JS"],
            ["huggingface_hub.InferenceClient", "Single client for all HF Inference API calls (ASR, NLP, chat)"],
            ["gtts.gTTS", "Google Text-to-Speech — converts reply text to MP3"],
            ["pandas", "Reads metrics_log.csv for the dashboard"],
            ["rouge_score", "Computes ROUGE-1 F1 to measure summary quality"],
            ["mlflow", "Optional LLMOps experiment tracking (app still works without it)"],
            ["os, time, csv, datetime, tempfile", "Standard library — file paths, timing, CSV writing, temp MP3 storage"],
        ],
        col_widths=[2.5, 4.2],
    )
    para(doc,
        "The mlflow import is wrapped in a try/except so the app still runs even "
        "if mlflow is not installed. This is a common pattern for optional dependencies."
    )

    # 3.2 Configuration
    heading(doc, "3.2  Configuration Constants (Lines 48–70)", 2)
    para(doc,
        "All model names and tunable parameters are declared as constants at the top, "
        "so they are easy to change without hunting through the code:"
    )
    table(doc,
        headers=["Constant", "Value", "Purpose"],
        rows=[
            ["HF_TOKEN", "Read from env var", "Authenticates all HF API calls"],
            ["ASR_MODEL", "openai/whisper-large-v3", "Speech-to-text model"],
            ["SENTIMENT_MODEL", "distilbert-base-uncased-finetuned-sst-2-english", "Sentiment classifier"],
            ["ZEROSHOT_MODEL", "facebook/bart-large-mnli", "Zero-shot intent classifier"],
            ["NER_MODEL", "dslim/bert-base-NER", "Named entity recogniser"],
            ["SUMMARY_MODEL", "facebook/bart-large-cnn", "Text summariser"],
            ["CHAT_MODEL", "Qwen/Qwen2.5-7B-Instruct", "LLM for reply generation"],
            ["FINETUNED_DIR", "./intent-model", "Path to our fine-tuned DistilBERT"],
            ["CANDIDATE_INTENTS", "7 strings", "Labels used by the zero-shot model"],
            ["EST_COST_PER_1K_TOKENS", "0.0002", "Illustrative cost for LLMOps cost tracking"],
        ],
        col_widths=[2.0, 2.3, 2.4],
    )

    # 3.3 LLMOps Logger
    heading(doc, "3.3  LLMOps Metrics Logger (Lines 72–108)", 2)
    para(doc,
        "Two small functions power all metrics logging:"
    )

    heading(doc, "log_metric(subtask, model, latency, confidence, tokens, extra)", 3)
    para(doc,
        "Called after every API call. It:"
    )
    bullet(doc, "Appends one row to metrics_log.csv (creates the file with headers if it doesn't exist)")
    bullet(doc, "Computes estimated cost: tokens / 1000 × EST_COST_PER_1K_TOKENS")
    bullet(doc, "If an MLflow run is active, mirrors the same numbers with mlflow.log_metric()")
    para(doc,
        "The CSV approach means the dashboard works even if MLflow is not installed. "
        "MLflow is the 'bonus' layer for industry-standard experiment tracking."
    )

    heading(doc, "timed(subtask, model_name, fn, *args, **kwargs)", 3)
    para(doc,
        "A thin wrapper that records wall-clock time around any function call:"
    )
    code_block(doc, "result, latency = timed('ASR', ASR_MODEL, client.automatic_speech_recognition, audio_path, model=ASR_MODEL)")
    para(doc,
        "*args and **kwargs let it wrap any function with any signature. "
        "It returns (result, latency) so the caller can both use the result "
        "and pass latency to log_metric()."
    )

    # 3.4 Sub-tasks
    heading(doc, "3.4  The 7 Sub-task Functions (Lines 111–245)", 2)

    # STT
    heading(doc, "Sub-task 1 — speech_to_text(audio_path)", 3)
    code_block(doc, "result = client.automatic_speech_recognition(audio_path, model=ASR_MODEL)")
    para(doc,
        "Takes an audio file path (microphone or uploaded file). Sends the audio "
        "to Whisper-large-v3 on the HF Inference API. Returns the transcript string. "
        "If no audio is provided (user typed instead), returns an empty string so "
        "the pipeline falls through to the typed text."
    )
    para(doc, "Key detail: the result object has a .text attribute. We guard with hasattr() "
         "in case the API response format changes.", italic=True)

    # TTS
    heading(doc, "Sub-task 2 — text_to_speech(text)", 3)
    code_block(doc, "tts = gTTS(text=text[:500], lang='en')\ntts.save(out_path)")
    para(doc,
        "Takes the AI-generated reply string. Creates a gTTS object (calls the Google "
        "TTS API), saves the MP3 to the system temp folder, returns the file path. "
        "Gradio's Audio component reads the path and plays it in the browser."
    )
    para(doc, "text[:500] — we cap at 500 characters to keep the audio short and avoid "
         "very long TTS API calls.", italic=True)

    # Sentiment
    heading(doc, "Sub-task 3 — analyze_sentiment(text)", 3)
    code_block(doc, "result = client.text_classification(text, model=SENTIMENT_MODEL)")
    para(doc,
        "Sends the transcript to DistilBERT-SST-2. Returns a list of "
        "ClassificationLabel objects ordered by score. We extract:"
    )
    bullet(doc, "scores dict  {label: score}  — used by Gradio's Label widget to show a bar chart")
    bullet(doc, "top.label   — 'POSITIVE' or 'NEGATIVE'  — injected into the LLM prompt")
    bullet(doc, "top.score   — confidence  — logged as a metric")

    # Intent zero-shot
    heading(doc, "Sub-task 4a — classify_intent_zeroshot(text)", 3)
    code_block(doc,
        "result = client.zero_shot_classification(\n"
        "    text,\n"
        "    candidate_labels=CANDIDATE_INTENTS,\n"
        "    model=ZEROSHOT_MODEL\n"
        ")"
    )
    para(doc,
        "Zero-shot classification means the model was never explicitly trained on "
        "these 7 labels. Instead, BART-large-MNLI (a Natural Language Inference model) "
        "is repurposed: it scores whether each label is 'entailed' by the text. "
        "No training data or JSONL file is needed — just the list of candidate intents."
    )
    para(doc, "Trade-off: convenient but confidence scores tend to be lower (~0.5) because "
         "the model is doing a task it wasn't originally designed for.", italic=True)

    # Intent fine-tuned
    heading(doc, "Sub-task 4b — classify_intent_finetuned(text)", 3)
    code_block(doc,
        "_finetuned_pipe = pipeline('text-classification', model=FINETUNED_DIR)\n"
        "out = _finetuned_pipe(text[:512])[0]"
    )
    para(doc,
        "Loads our fine-tuned DistilBERT from ./intent-model on first call and caches "
        "it in the global _finetuned_pipe variable (lazy loading — avoids loading the "
        "model at startup). The model was explicitly trained on 27 real support intents "
        "from the Bitext dataset, so its predictions are more domain-accurate."
    )
    para(doc,
        "The function gracefully falls back to a 'model not found' message if the "
        "./intent-model folder is missing, so the app still runs."
    )
    para(doc, "text[:512] — DistilBERT has a 512-token context limit. Truncating here "
         "prevents errors on very long inputs.", italic=True)

    # NER
    heading(doc, "Sub-task 5 — extract_entities(text)", 3)
    code_block(doc, "result = client.token_classification(text, model=NER_MODEL)")
    para(doc,
        "Token classification (NER) labels each word/subword token with an entity type: "
        "PER (person), ORG (organisation), LOC (location), MISC. "
        "The result is a list of token spans with start/end character offsets. "
        "Gradio's HighlightedText component uses these offsets to highlight the "
        "entities in the original transcript with colour coding."
    )
    para(doc, "e.g. 'Priya' → PER, 'Chennai' → LOC — gives the human agent instant "
         "structured context without reading the full complaint.", italic=True)

    # Summarization
    heading(doc, "Sub-task 6 — summarize(text)", 3)
    code_block(doc,
        "if len(text.split()) < 15:\n"
        "    return text   # too short to summarise\n"
        "result = client.summarization(text, model=SUMMARY_MODEL)\n"
        "rouge1 = _rouge.score(text, summary)['rouge1'].fmeasure"
    )
    para(doc,
        "Sends the transcript to BART-large-CNN (trained on CNN/DailyMail news "
        "articles — generalises well to complaint text). After receiving the summary, "
        "ROUGE-1 F1 is computed locally (no API call) by comparing word overlap "
        "between the summary and the original. This score is logged as the "
        "'confidence' column for the summarization row in the CSV."
    )
    para(doc, "ROUGE-1 F1 measures unigram (single word) overlap. A score of 0.55 means "
         "~55% of the summary words appeared in the original, balanced with recall. "
         "It's a proxy for faithfulness, not quality.", italic=True)

    # Response generation
    heading(doc, "Sub-task 7 — generate_response(complaint, sentiment_label, intent)", 3)
    code_block(doc,
        "resp = client.chat_completion(\n"
        "    messages=[{'role': 'user', 'content': prompt}],\n"
        "    model=CHAT_MODEL,\n"
        "    max_tokens=250,\n"
        ")"
    )
    para(doc,
        "The prompt is constructed by injecting the detected sentiment and intent "
        "into a system-role description:"
    )
    code_block(doc,
        "\"You are a polite customer-support agent...\n"
        " The customer's message (sentiment: NEGATIVE, detected intent: refund request) is:\n"
        " <complaint text>\n"
        " Write a short, empathetic reply (max 120 words) with a concrete next step.\""
    )
    para(doc,
        "This is called 'prompt grounding' — using the outputs of earlier pipeline "
        "steps to make the LLM's reply more accurate and appropriate to the situation. "
        "The response is generated by Qwen2.5-7B via HF's OpenAI-compatible chat API "
        "(same interface as the OpenAI SDK, just a different base URL and model name)."
    )
    para(doc, "max_tokens=250 keeps replies concise. resp.usage.total_tokens is logged "
         "for cost tracking.", italic=True)

    # 3.5 Pipeline orchestrator
    heading(doc, "3.5  run_pipeline(audio, typed_text) — The Orchestrator (Lines 251–308)", 2)
    para(doc,
        "This is the central function that wires all 7 sub-tasks together. "
        "It is a Python generator (uses yield instead of return), which is how "
        "Gradio achieves streaming — the UI updates after every sub-task completes "
        "rather than waiting for all 7 to finish."
    )
    para(doc, "How the generator pattern works:")
    bullet(doc, "The function yields a tuple of all output values after each step")
    bullet(doc, "Gradio calls next() on the generator repeatedly and updates the UI each time")
    bullet(doc, "The user sees results filling in live: transcript first, then sentiment, etc.")
    para(doc, "Pipeline logic:")
    bullet(doc, "If audio is provided → call speech_to_text() for the transcript")
    bullet(doc, "If only text is provided → use it directly as the transcript")
    bullet(doc, "If neither → yield an error message and return")
    bullet(doc, "Run sub-tasks 2–7 in order, yielding status updates between each")
    bullet(doc, "Wrap in try/except to catch any API error and display it gracefully")
    bullet(doc, "If MLflow is on: start a run before the pipeline, end it after (or on error)")

    code_block(doc,
        "def run_pipeline(audio, typed_text):\n"
        "    ...yield snap('Step 1/7 — transcribing...')     # updates UI immediately\n"
        "    transcript = speech_to_text(audio)\n"
        "    out['transcript'] = transcript\n"
        "    yield snap('Step 2/7 — analysing sentiment...')  # next UI update\n"
        "    ..."
    )

    # 3.6 Feedback
    heading(doc, "3.6  record_feedback(feedback) (Lines 311–317)", 2)
    para(doc,
        "Called when the user clicks the 👍 / 👎 radio button. Logs a 1.0 or 0.0 "
        "confidence score to the metrics CSV and to MLflow. This is the 'user "
        "satisfaction' LLMOps metric — the only human-in-the-loop signal in the pipeline."
    )

    # 3.7 Dashboard
    heading(doc, "3.7  load_dashboard() — LLMOps Dashboard (Lines 320–353)", 2)
    para(doc,
        "Reads metrics_log.csv with pandas and computes aggregate statistics "
        "grouped by sub-task:"
    )
    table(doc,
        headers=["Aggregation", "Purpose"],
        rows=[
            ["count", "How many times each sub-task has been called"],
            ["mean(latency_sec)", "Average response time — identifies slow sub-tasks"],
            ["mean(confidence)", "Average quality signal — drops in confidence indicate degradation"],
            ["sum(tokens)", "Total token consumption — drives cost estimates"],
            ["sum(est_cost_usd)", "Total estimated spend across all runs"],
        ],
        col_widths=[2.0, 4.7],
    )
    para(doc,
        "The function returns four values: the summary dataframe (for the table), "
        "a KPI markdown string, and two filtered dataframes for the two bar charts. "
        "Gradio's BarPlot component renders the charts directly from pandas DataFrames."
    )

    # 3.8 Gradio UI
    heading(doc, "3.8  Gradio UI Definition (Lines 356–433)", 2)
    para(doc,
        "The UI is defined inside a gr.Blocks() context manager. Key components:"
    )
    table(doc,
        headers=["Component", "Type", "Purpose"],
        rows=[
            ["audio_in", "gr.Audio", "Microphone or file upload for voice complaint"],
            ["text_in", "gr.Textbox", "Alternative typed complaint input"],
            ["gr.Examples", "—", "One-click sample complaints for demo"],
            ["run_btn", "gr.Button", "Triggers run_pipeline() on click"],
            ["transcript_out", "gr.Textbox", "Shows Whisper transcript"],
            ["sentiment_out", "gr.Label", "Bar chart of POSITIVE/NEGATIVE scores"],
            ["intent_zs_out", "gr.Label", "Bar chart of top-3 zero-shot intents"],
            ["intent_ft_out", "gr.Label", "Shows fine-tuned model prediction + score"],
            ["ner_out", "gr.HighlightedText", "Highlights named entities in colour"],
            ["summary_out", "gr.Textbox", "Shows ticket summary"],
            ["reply_out", "gr.Textbox", "Shows AI-generated reply"],
            ["reply_audio_out", "gr.Audio", "Plays spoken reply (autoplay=True)"],
            ["fb (Radio)", "gr.Radio", "👍/👎 feedback → record_feedback()"],
            ["lat_plot / conf_plot", "gr.BarPlot", "Dashboard charts from metrics_log.csv"],
        ],
        col_widths=[1.6, 1.2, 3.9],
    )
    para(doc,
        "The run_btn.click() call wires the button to run_pipeline() and maps its "
        "outputs to the display components. Because run_pipeline is a generator, "
        "Gradio automatically streams the updates. The Dashboard tab uses "
        "dash_tab.select() to auto-refresh whenever the tab is opened."
    )

    doc.add_page_break()

    # ── finetune_intent.py deep dive ──────────────────────────────────────────
    heading(doc, "4. finetune_intent.py — Deep Dive", 1)
    para(doc,
        "This is a standalone script — you run it once to produce the ./intent-model "
        "folder, then never need to run it again (unless you want to retrain)."
    )

    heading(doc, "4.1  Why Fine-tune? Zero-shot vs Fine-tuned", 2)
    table(doc,
        headers=["", "Zero-shot (BART-MNLI)", "Fine-tuned (DistilBERT)"],
        rows=[
            ["Training data needed", "None", "Bitext dataset (~8k samples)"],
            ["Setup time", "Instant", "~100 seconds"],
            ["Accuracy on support intents", "~40–60% confidence", "97%+ accuracy"],
            ["Number of intents", "7 (our custom list)", "27 (real Bitext intents)"],
            ["Inference speed", "Slow (large BART, API call)", "Very fast (small DistilBERT, local)"],
            ["Use case", "Quick prototype / unknown domains", "Production — known domain"],
        ],
        col_widths=[1.8, 2.2, 2.7],
    )

    heading(doc, "4.2  Dataset — Bitext Customer Support", 2)
    para(doc,
        "The Bitext Customer Support LLM Chatbot Training Dataset contains ~27,000 "
        "labelled utterances across 27 real customer support intents such as:"
    )
    bullet(doc, "cancel_order, get_refund, track_order, delivery_period")
    bullet(doc, "payment_issue, check_invoice, contact_human_agent")
    bullet(doc, "create_account, delete_account, recover_password")
    para(doc,
        "It is loaded directly from HuggingFace Hub using load_dataset() — "
        "no manual download, no JSONL file needed. The datasets library downloads "
        "it in Apache Arrow format (columnar binary), which the HF Trainer reads natively."
    )
    para(doc, "Why not JSONL? JSONL is required only for OpenAI's fine-tuning API. "
         "HuggingFace Trainer uses the datasets library's Arrow format internally. "
         "You never need to create or touch a JSONL file for this approach.", italic=True)

    heading(doc, "4.3  Script Walkthrough — Step by Step", 2)

    heading(doc, "Step 1: Load and Subsample (Lines 45–57)", 3)
    code_block(doc,
        "ds = load_dataset(DATASET, split='train').shuffle(seed=42)\n"
        "ds = ds.select(range(8000))  # use 8k of 27k for speed\n"
        "labels = sorted(set(ds['intent']))  # 27 unique intent strings\n"
        "label2id = {l: i for i, l in enumerate(labels)}  # 'cancel_order' → 0 etc.\n"
        "ds = ds.map(lambda x: {'label': label2id[x['intent']]})  # add int label column\n"
        "ds = ds.train_test_split(test_size=0.15, seed=42)  # 6800 train / 1200 test"
    )
    para(doc,
        "shuffle(seed=42) ensures reproducibility. seed=42 is a convention — any "
        "fixed number gives the same shuffle every run. The 15% test split gives "
        "~1,200 samples for evaluation, never seen during training."
    )

    heading(doc, "Step 2: Tokenize (Lines 59–67)", 3)
    code_block(doc,
        "tok = AutoTokenizer.from_pretrained('distilbert-base-uncased')\n"
        "def tokenize(batch):\n"
        "    return tok(batch['instruction'], truncation=True, max_length=128)\n"
        "ds = ds.map(tokenize, batched=True)"
    )
    para(doc,
        "Tokenization converts raw text into integer token IDs that the model understands. "
        "The Bitext dataset stores the utterance in a column called 'instruction'. "
        "truncation=True and max_length=128 ensure no input exceeds DistilBERT's limit. "
        "batched=True processes many examples at once for speed."
    )

    heading(doc, "Step 3: Load the Base Model (Lines 70–73)", 3)
    code_block(doc,
        "model = AutoModelForSequenceClassification.from_pretrained(\n"
        "    'distilbert-base-uncased',\n"
        "    num_labels=27,\n"
        "    id2label=id2label,\n"
        "    label2id=label2id,\n"
        ")"
    )
    para(doc,
        "AutoModelForSequenceClassification takes the standard DistilBERT (pre-trained "
        "on English Wikipedia + BookCorpus for language understanding) and adds a "
        "fresh classification head — a single linear layer that maps the 768-dim "
        "hidden state to 27 output classes. The pre-trained weights are UNEXPECTED "
        "for this new head (you'll see that in the load report), which is normal."
    )

    heading(doc, "Step 4: Define Metrics (Lines 75–86)", 3)
    code_block(doc,
        "accuracy = evaluate.load('accuracy')\n"
        "f1 = evaluate.load('f1')\n"
        "def compute_metrics(eval_pred):\n"
        "    preds = np.argmax(logits, axis=-1)  # pick highest-score class\n"
        "    return {\n"
        "        'accuracy': accuracy.compute(...)['accuracy'],\n"
        "        'macro_f1': f1.compute(..., average='macro')['f1'],\n"
        "    }"
    )
    para(doc,
        "Accuracy = % of test samples predicted correctly. "
        "Macro-F1 = average F1 across all 27 classes, giving equal weight to each "
        "intent regardless of how many samples it has. Macro-F1 is more informative "
        "than accuracy when classes are imbalanced."
    )

    heading(doc, "Step 5: Training Arguments (Lines 88–99)", 3)
    table(doc,
        headers=["Argument", "Value", "Why"],
        rows=[
            ["num_train_epochs", "2", "2 passes through the data; more gives marginal gain but takes longer"],
            ["per_device_train_batch_size", "32", "32 samples processed together; larger = faster but needs more RAM"],
            ["learning_rate", "2e-5", "Standard fine-tuning rate for BERT-family models (too high = forgetting pre-training)"],
            ["eval_strategy", "epoch", "Evaluate on the test set after every epoch to track progress"],
            ["save_strategy", "no", "Don't save checkpoints mid-training (we save the final model explicitly)"],
            ["report_to", "none", "Disable WandB/TensorBoard; we use MLflow instead"],
        ],
        col_widths=[2.0, 1.2, 3.5],
    )

    heading(doc, "Step 6: Train, Evaluate, Save (Lines 101–148)", 3)
    code_block(doc,
        "trainer = Trainer(model, args, train_dataset, eval_dataset, ...)\n"
        "trainer.train()\n"
        "metrics = trainer.evaluate()\n"
        "print(f'Accuracy: {metrics[\"eval_accuracy\"]:.4f}')  # 97.07%\n"
        "print(f'Macro F1: {metrics[\"eval_macro_f1\"]:.4f}')  # 97.07%\n"
        "trainer.save_model('./intent-model')\n"
        "tok.save_pretrained('./intent-model')"
    )
    para(doc,
        "The Trainer handles the training loop, gradient descent, mixed precision, "
        "and evaluation automatically. After training, both the model weights and the "
        "tokenizer are saved to ./intent-model — the tokenizer must be saved alongside "
        "the model because the app needs it to tokenize text at inference time."
    )
    para(doc,
        "The final sanity check at the bottom tests 3 sample sentences and prints "
        "their predicted intents. This is a quick visual check that the model is "
        "making sensible predictions before the training script exits."
    )

    doc.add_page_break()

    # ── Data Flow ─────────────────────────────────────────────────────────────
    heading(doc, "5. End-to-End Data Flow", 1)
    para(doc,
        "A single complaint flows through the system as follows:"
    )
    table(doc,
        headers=["Stage", "Input", "Process", "Output", "Logged Metric"],
        rows=[
            ["1", "Audio file / text", "Whisper-large-v3 (HF API)", "Transcript string", "Latency, word count"],
            ["2", "Transcript", "DistilBERT-SST-2 (HF API)", "POSITIVE/NEGATIVE + score", "Latency, confidence"],
            ["3", "Transcript", "BART-large-MNLI (HF API)", "Intent label + score (zero-shot)", "Latency, confidence"],
            ["4", "Transcript", "Fine-tuned DistilBERT (local)", "Intent label + score", "Latency, confidence"],
            ["5", "Transcript", "BERT-base-NER (HF API)", "[(entity, type, start, end)]", "Latency, avg confidence"],
            ["6", "Transcript", "BART-large-CNN (HF API)", "Summary text + ROUGE-1 F1", "Latency, ROUGE-1, tokens"],
            ["7", "Transcript + sentiment + intent", "Qwen2.5-7B (HF chat API)", "Reply text (≤120 words)", "Latency, tokens, cost"],
            ["8", "Reply text", "gTTS API", "MP3 audio file", "Latency, word count"],
            ["9", "User click", "record_feedback()", "1.0 or 0.0 score", "User satisfaction"],
        ],
        col_widths=[0.5, 1.4, 1.8, 1.7, 1.3],
    )

    # ── Frequently Asked Questions ────────────────────────────────────────────
    heading(doc, "6. Frequently Asked Questions", 1)

    faqs = [
        ("Do I need to provide a JSONL file for fine-tuning?",
         "No. The fine-tuning uses HuggingFace's datasets library which downloads "
         "the Bitext dataset directly from the HF Hub. JSONL is only needed for "
         "OpenAI's proprietary fine-tuning API — a completely different approach."),

        ("Where do I get the HF_TOKEN?",
         "Go to https://huggingface.co/settings/tokens, click 'New token', choose "
         "Read access, copy the hf_xxx... string. Each team member needs their own "
         "token. It's free. Set it with: export HF_TOKEN=hf_your_token"),

        ("The app says 'fine-tuned model not found'. What do I do?",
         "Run python3 finetune_intent.py from the voicedesk-ai folder. It takes "
         "~100 seconds on Apple Silicon. After it finishes, ./intent-model will "
         "exist and the app will load it automatically on next run."),

        ("Why is the fine-tuned model's confidence score low (~0.2) in the dashboard?",
         "The fine-tuned model has 27 intents, so probability mass is spread across "
         "more classes — lower individual scores are expected. A score of 0.2 with "
         "27 classes is actually decisive (random would be ~0.037). The accuracy "
         "(97%) is the right quality signal, not the raw confidence."),

        ("Can the app work without an HF token?",
         "Partially. The HF Inference API rate-limits unauthenticated requests "
         "heavily. The fine-tuned intent model and gTTS work without a token, but "
         "Whisper, sentiment, zero-shot intent, NER, summarization, and the LLM "
         "reply will fail or be very slow without one."),

        ("What if the MLflow dashboard is empty?",
         "Make sure to add --backend-store-uri sqlite:///mlflow.db to the mlflow ui "
         "command. Without it, MLflow looks in ./mlruns (empty) instead of mlflow.db."),

        ("Can I change the chat model (Qwen2.5-7B)?",
         "Yes — change CHAT_MODEL in app.py to any model listed at "
         "router.huggingface.co/v1/models. The API call uses OpenAI-compatible "
         "chat_completion, so any HF-served chat model works without changing other code."),
    ]

    for q, a in faqs:
        para(doc, f"Q: {q}", bold=True)
        para(doc, f"A: {a}")
        doc.add_paragraph()

    doc.save("VoiceDesk_AI_Code_Explainer.docx")
    print("VoiceDesk_AI_Code_Explainer.docx generated successfully.")


if __name__ == "__main__":
    build()
